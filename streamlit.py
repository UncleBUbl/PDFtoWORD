import streamlit as st
import google.generativeai as genai
from docx import Document
from io import BytesIO

# --- CONFIGURATION ---
st.set_page_config(layout="centered", page_title="DocuGenius Pro", page_icon="⚡")

# --- CSS STYLING ---
st.markdown("""
<style>
    .stApp { background-color: #f8f9fa; }
    
    /* Editor Styling */
    .stTextArea textarea { 
        font-family: 'Courier New', monospace; 
        font-size: 15px; 
        line-height: 1.6;
        background-color: #ffffff;
        color: #333;
        border: 1px solid #ddd;
    }
    
    .main-header { 
        font-size: 2.5rem; 
        font-weight: 800; 
        color: #1a73e8; 
        text-align: center; 
        margin-bottom: 1rem;
    }
    
    .sub-header {
        text-align: center;
        color: #666;
        margin-bottom: 2rem;
    }
</style>
""", unsafe_allow_html=True)

# --- HELPER: Advanced Markdown Tables to Docx ---
def create_table_from_markdown(doc, table_lines):
    """Parses Markdown table lines and builds a real Word table."""
    try:
        # Filter out separator lines (like |---|---|)
        rows = [row for row in table_lines if '---' not in row]
        if not rows: return

        # Calculate columns based on header
        header_cells = [c.strip() for c in rows[0].split('|') if c.strip() != '']
        num_cols = len(header_cells)
        
        # Create Table
        table = doc.add_table(rows=len(rows), cols=num_cols)
        table.style = 'Table Grid'
        
        for i, row_text in enumerate(rows):
            cells_text = [c.strip() for c in row_text.split('|') if c.strip() != '']
            for j, text in enumerate(cells_text):
                if j < num_cols:
                    table.cell(i, j).text = text
    except:
        # Fallback if parsing fails
        for line in table_lines:
            doc.add_paragraph(line)

def generate_docx(text_content):
    doc = Document()
    lines = text_content.split('\n')
    table_buffer = []
    
    for line in lines:
        stripped = line.strip()
        
        # Detect Table Rows
        if stripped.startswith('|') and stripped.endswith('|'):
            table_buffer.append(stripped)
        else:
            # Write buffered table if we just finished one
            if table_buffer:
                create_table_from_markdown(doc, table_buffer)
                table_buffer = [] 
            
            # Normal Text Processing
            if line.startswith('# '): doc.add_heading(line[2:], level=1)
            elif line.startswith('## '): doc.add_heading(line[2:], level=2)
            elif line.startswith('### '): doc.add_heading(line[2:], level=3)
            elif line.startswith('* ') or line.startswith('- '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif stripped:
                doc.add_paragraph(line)
    
    # Check for table at end of file
    if table_buffer:
        create_table_from_markdown(doc, table_buffer)

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- MAIN APP ---
def main():
    # 1. Sidebar
    with st.sidebar:
        st.title("⚙️ Settings")
        try:
            api_key = st.secrets["GOOGLE_API_KEY"]
            st.success("API Key Active")
        except:
            api_key = st.text_input("Gemini API Key", type="password")
            
        st.info("Using Model: **Gemini 3 Pro Preview**")

    # 2. Header
    st.markdown('<div class="main-header">DocuGenius Pro</div>', unsafe_allow_html=True)
    st.markdown('<div class="sub-header">Convert PDFs to Word (with Tables)</div>', unsafe_allow_html=True)

    # 3. Session State
    if 'generated_content' not in st.session_state:
        st.session_state.generated_content = ""

    # 4. Upload View
    if not st.session_state.generated_content:
        uploaded_file = st.file_uploader("Upload PDF", type=['pdf'])
        
        if uploaded_file and st.button("🚀 Convert PDF", type="primary", use_container_width=True):
            if not api_key:
                st.error("Please enter an API Key in the sidebar.")
            else:
                genai.configure(api_key=api_key)
                
                with st.spinner("Gemini 3 Pro is processing..."):
                    try:
                        # --- STRICTLY USING GEMINI 3 PRO PREVIEW ---
                        model = genai.GenerativeModel('gemini-3-pro-preview')
                        
                        prompt = """
                        You are a document conversion engine. Extract all content from this PDF.
                        
                        STRICT RULES:
                        1. Output clean Markdown.
                        2. TABLES: specific attention to tables. You MUST format them as Markdown tables.
                           Example:
                           | Header 1 | Header 2 |
                           |---|---|
                           | Data 1   | Data 2   |
                        3. Do not miss any rows.
                        4. Do not include conversational text like "Here is the output".
                        """

                        response = model.generate_content([
                            {'mime_type': 'application/pdf', 'data': uploaded_file.getvalue()},
                            prompt
                        ])
                        
                        st.session_state.generated_content = response.text
                        st.rerun()
                    except Exception as e:
                        st.error(f"Error: {e}")
                        if "404" in str(e):
                             st.warning("Model not found. Check your API key or try 'gemini-1.5-pro' if Gemini 3 is not enabled.")

    # 5. Editor View (Full Width)
    else:
        st.success("✅ Conversion Complete! Edit below.")
        
        edited_text = st.text_area(
            "📝 Edit your document (Markdown format):", 
            value=st.session_state.generated_content, 
            height=600
        )
        st.session_state.generated_content = edited_text
        
        st.divider()
        
        col1, col2 = st.columns([1, 1])
        
        with col1:
             if st.button("🔄 Start New File", use_container_width=True):
                st.session_state.generated_content = ""
                st.rerun()
        
        with col2:
            docx_data = generate_docx(st.session_state.generated_content)
            st.download_button(
                label="⬇️ Download Word Doc", 
                data=docx_data, 
                file_name="converted_gemini3.docx", 
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary",
                use_container_width=True
            )

if __name__ == "__main__":
    main()
